#!/usr/bin/env python3
"""
nucleo_despachadora/indexar_corpus.py
Portado de: G:\\Meu Drive\\...\\Skill despachadora de Docs\\indexar_corpus.py

Adaptações aplicadas:
  • SKILL_ROOT hardcoded → lido de segredos.env via Segredos (get("corpus")["path"])
                           resolvido DENTRO de main(), não no import
  • corpus_index.json    → gravado em automacoes/despachadora/ (_OUTPUT_FILE)
                           onde despachadora.py o encontra
  • Execução standalone  → python indexar_corpus.py  (fora do painel)

Uso:
    cd automacoes/despachadora/nucleo_despachadora
    python indexar_corpus.py

Pré-condição: CORPUS_PATH preenchido em segredos.env (raiz do projeto).

Indexa todos os documentos da Skill despachadora de Docs.
Saída: corpus_index.json em automacoes/despachadora/.

Extração por tipo:
  PDF   -> PyMuPDF -> pypdf -> error:"pdf_imagem_sem_ocr"
  DOCX  -> python-docx
  DOC   -> win32com Word.Application (batch, instância única)
  XLSX/XLS -> openpyxl -> error:"xlsx_falhou"
  TXT/MD/HTML/HTM -> leitura direta (utf-8 -> latin-1 -> cp1252)

Exclusões:
  Pastas : P1/Holerites
           P1/Holerites Bonus DEJEM
           P4/endereços de ip
  Extensões : .pfx .cer .rar .mp4 .jpg .jpeg .png .pptx .ppt
  Arquivo   : despachadora.py (raiz da skill)

Modo incremental:
  Se corpus_index.json já existir, carrega as entradas presentes e processa
  apenas os arquivos cujo par (section, arquivo) ainda não esteja indexado.
  Entradas existentes são preservadas intactas.
"""

from __future__ import annotations

import os
import json
import sys
from pathlib import Path

# Forçar UTF-8 no stdout/stderr
if hasattr(sys.stdout, "reconfigure"):
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
if hasattr(sys.stderr, "reconfigure"):
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")

# Garante que nucleo/ seja importável mesmo rodando standalone
_NUCLEO_DESP = Path(__file__).resolve().parent          # nucleo_despachadora/
_DESP_DIR    = _NUCLEO_DESP.parent                      # automacoes/despachadora/
_ROOT        = _DESP_DIR.parent.parent                  # raiz do projeto
for _p in (str(_ROOT), str(_NUCLEO_DESP)):
    if _p not in sys.path:
        sys.path.insert(0, _p)

from nucleo.segredos import Segredos

# corpus_index.json fica em automacoes/despachadora/ (mesmo nível de executar.py)
# _SKILL_ROOT é resolvido em main() para não executar no import
_OUTPUT_FILE = _DESP_DIR / "corpus_index.json"

# ── Constantes ────────────────────────────────────────────────────────────────

EXCLUDED_DIR_PREFIXES = [
    Path("P1") / "Holerites",
    Path("P1") / "Holerites Bonus DEJEM",
    Path("P4") / "endereços de ip",
]

EXCLUDED_FILE_NAMES = {
    "despachadora.py",
    "indexar_corpus.py",
    "corpus_index.json",
}

EXCLUDED_EXTENSIONS = {
    ".pfx", ".cer", ".rar",
    ".mp4",
    ".jpg", ".jpeg", ".png",
    ".pptx", ".ppt",
}


# ── Helpers de caminho ────────────────────────────────────────────────────────

def is_excluded_dir(rel_from_root: Path) -> bool:
    """True se rel_from_root é igual a, ou filho de, um dos prefixos excluídos."""
    for prefix in EXCLUDED_DIR_PREFIXES:
        try:
            rel_from_root.relative_to(prefix)
            return True
        except ValueError:
            pass
    return False


def to_posix(p: Path) -> str:
    """Converte separadores para '/' independente de OS."""
    return str(p).replace(os.sep, "/")


def incremental_key(section: str, arquivo: str) -> str:
    return f"{section}/{arquivo}" if section else arquivo


# ── Extratores ────────────────────────────────────────────────────────────────

def extract_pdf(path: Path):
    """PyMuPDF -> pypdf -> ('', 'pdf_imagem_sem_ocr')."""
    try:
        import fitz
        doc = fitz.open(str(path))
        text = "".join(page.get_text() for page in doc)
        doc.close()
        if text.strip():
            return text, None
    except Exception:
        pass

    try:
        from pypdf import PdfReader
        reader = PdfReader(str(path))
        text = "".join(page.extract_text() or "" for page in reader.pages)
        if text.strip():
            return text, None
    except Exception:
        pass

    return "", "pdf_imagem_sem_ocr"


def extract_docx(path: Path):
    try:
        from docx import Document
        doc = Document(str(path))
        parts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text)
        return "\n".join(parts), None
    except Exception:
        return "", "docx_falhou"


def extract_xlsx(path: Path):
    try:
        import openpyxl
        wb = openpyxl.load_workbook(str(path), read_only=True, data_only=True)
        lines = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            for row in ws.iter_rows(values_only=True):
                row_text = "\t".join("" if c is None else str(c) for c in row)
                if row_text.strip():
                    lines.append(row_text)
        wb.close()
        return "\n".join(lines), None
    except Exception:
        return "", "xlsx_falhou"


def extract_text_file(path: Path):
    for enc in ("utf-8", "latin-1", "cp1252"):
        try:
            return path.read_text(encoding=enc), None
        except Exception:
            continue
    return "", "txt_encoding_falhou"


def extract_generic(path: Path):
    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_pdf(path)
    if ext == ".docx":
        return extract_docx(path)
    if ext in (".xlsx", ".xls"):
        return extract_xlsx(path)
    if ext in (".txt", ".md", ".html", ".htm"):
        return extract_text_file(path)
    return "", "tipo_nao_suportado"


# ── Batch COM para .doc ───────────────────────────────────────────────────────

def batch_extract_doc(doc_paths: list) -> dict:
    """
    Processa todos os .doc com uma única instância do Word.Application.
    Retorna dict: str(path.resolve()) -> (texto, error_ou_None).
    """
    results = {}
    if not doc_paths:
        return results

    word = None
    try:
        import pythoncom
        import win32com.client
    except ImportError:
        for fpath in doc_paths:
            results[str(fpath.resolve())] = ("", "win32com_indisponivel")
        return results

    pythoncom.CoInitialize()
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False

        for fpath in doc_paths:
            abs_str = str(fpath.resolve())
            try:
                doc = word.Documents.Open(
                    abs_str,
                    ReadOnly=True,
                    AddToRecentFiles=False,
                    ConfirmConversions=False,
                )
                text = doc.Content.Text
                doc.Close(False)
                results[abs_str] = (text, None)
            except Exception:
                results[abs_str] = ("", "doc_com_falhou")

    except Exception:
        for fpath in doc_paths:
            abs_str = str(fpath.resolve())
            if abs_str not in results:
                results[abs_str] = ("", "doc_com_falhou")
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        pythoncom.CoUninitialize()

    return results


# ── Coleta de arquivos ────────────────────────────────────────────────────────

def collect_files(root: Path) -> list:
    """
    Retorna lista de (section: str, fpath: Path).
    section = nome do subdiretório de primeiro nível; '' para arquivos na raiz.
    Respeita todas as exclusões de diretório, extensão e nome.
    """
    result = []

    for entry in sorted(root.iterdir()):

        if entry.is_file():
            if entry.name in EXCLUDED_FILE_NAMES:
                continue
            if entry.name.startswith("~$"):
                continue
            if entry.suffix.lower() in EXCLUDED_EXTENSIONS:
                continue
            result.append(("", entry))

        elif entry.is_dir():
            section = entry.name
            for dirpath_str, dirnames, filenames in os.walk(str(entry), topdown=True):
                dp = Path(dirpath_str)
                rel_dp = dp.relative_to(root)

                if is_excluded_dir(rel_dp):
                    dirnames.clear()
                    continue

                dirnames[:] = sorted(
                    d for d in dirnames
                    if not is_excluded_dir(rel_dp / d)
                )

                for fname in sorted(filenames):
                    if fname in EXCLUDED_FILE_NAMES:
                        continue
                    if fname.startswith("~$"):
                        continue
                    fp = dp / fname
                    if fp.suffix.lower() in EXCLUDED_EXTENSIONS:
                        continue
                    result.append((section, fp))

    return result


def aplicar_curadoria(entries: list[dict], curadoria_path: Path) -> tuple[list[dict], dict]:
    """
    Aplica as regras do manifesto de curadoria (exclusão, reclassificação, rebaixamento)
    sobre os chunks de documentos.
    Retorna a lista filtrada/modificada de chunks e um dicionário com estatísticas.
    """
    stats = {
        "excluidos_chunks": 0, "excluidos_arquivos": set(),
        "reclassificados_chunks": 0, "reclassificados_arquivos": set(),
        "rebaixados_chunks": 0, "rebaixados_arquivos": set()
    }
    
    if not curadoria_path.exists():
        print(f"\nAviso: Manifesto de curadoria não localizado em {curadoria_path.name}. Nenhuma curadoria aplicada.")
        return entries, stats

    try:
        rules = json.loads(curadoria_path.read_text(encoding="utf-8"))
    except Exception as e:
        print(f"\nErro ao ler manifesto de curadoria: {e}. Nenhuma curadoria aplicada.")
        return entries, stats

    excluir_rules = rules.get("excluir_do_indice", [])
    reclassificar_rules = rules.get("reclassificar", [])
    rebaixar_rules = rules.get("rebaixar", [])

    def matches_rule(entry: dict, rule: dict) -> bool:
        pattern = rule.get("padrao_arquivo", "").replace("\\", "/")
        section = entry.get("section") or ""
        arquivo = (entry.get("arquivo") or "").replace("\\", "/")
        full_rel = f"{section}/{arquivo}" if section else arquivo
        
        # Check prefix/substring
        if pattern in full_rel:
            return True
        # Check base name
        if pattern in os.path.basename(arquivo):
            return True
        return False

    filtered_entries = []
    
    for entry in entries:
        section = entry.get("section") or ""
        arquivo = entry.get("arquivo") or ""
        file_key = f"{section}/{arquivo}" if section else arquivo
        
        # 1. Verificar exclusão
        is_excluido = False
        for rule in excluir_rules:
            if matches_rule(entry, rule):
                stats["excluidos_chunks"] += 1
                stats["excluidos_arquivos"].add(file_key)
                is_excluido = True
                break
                
        if is_excluido:
            continue

        # 2. Verificar reclassificação
        is_reclassificado = False
        for rule in reclassificar_rules:
            if matches_rule(entry, rule):
                natureza_nova = rule.get("natureza_nova")
                if entry.get("natureza") != natureza_nova:
                    entry["natureza"] = natureza_nova
                    stats["reclassificados_chunks"] += 1
                    stats["reclassificados_arquivos"].add(file_key)
                    is_reclassificado = True
                break

        # 3. Verificar rebaixamento
        is_rebaixado = False
        for rule in rebaixar_rules:
            if matches_rule(entry, rule):
                natureza_nova = rule.get("natureza_nova", "OUTRO")
                if entry.get("natureza") != natureza_nova:
                    entry["natureza"] = natureza_nova
                    stats["rebaixados_chunks"] += 1
                    stats["rebaixados_arquivos"].add(file_key)
                    is_rebaixado = True
                break

        filtered_entries.append(entry)

    return filtered_entries, stats


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    # _SKILL_ROOT resolvido aqui (não no import) para não bloquear despachadora.py
    _SKILL_ROOT = Path(Segredos().get("corpus")["path"])

    # ── Carregar índice existente (modo incremental) ──────────────────────────
    existing: dict = {}
    if _OUTPUT_FILE.exists():
        try:
            raw = json.loads(_OUTPUT_FILE.read_text(encoding="utf-8"))
            for entry in raw:
                k = incremental_key(entry.get("section", ""), entry["arquivo"])
                existing[k] = entry
            print(f"Modo incremental: {len(existing)} entradas carregadas de {_OUTPUT_FILE.name}")
        except Exception as e:
            print(f"Aviso: falha ao ler índice existente ({e}). Recriando do zero.")

    # ── Coletar todos os arquivos do corpus ───────────────────────────────────
    all_files = collect_files(_SKILL_ROOT)
    print(f"Arquivos encontrados no corpus: {len(all_files)}")

    # ── Separar novos vs já indexados ─────────────────────────────────────────
    to_process = []
    doc_batch  = []

    for section, fpath in all_files:
        section_root = _SKILL_ROOT / section if section else _SKILL_ROOT
        rel = to_posix(fpath.relative_to(section_root))
        key = incremental_key(section, rel)
        if key in existing:
            continue
        to_process.append((section, fpath, rel))
        if fpath.suffix.lower() in (".doc", ".rtf"):
            doc_batch.append(fpath)

    skipped = len(all_files) - len(to_process)
    print(f"  Já indexados (pulados)  : {skipped}")
    print(f"  Novos para processar    : {len(to_process)}")

    # ── Pré-processar todos os .doc em batch com uma única instância Word ─────
    doc_results = {}
    if doc_batch:
        print(f"\nAbrindo Word.Application para {len(doc_batch)} arquivo(s) .doc ...")
        doc_results = batch_extract_doc(doc_batch)
        print("Word.Application encerrado.\n")

    # ── Processar cada arquivo novo ───────────────────────────────────────────
    new_entries    = []
    section_counts = {}
    error_counts   = {}
    new_chars      = 0

    for section, fpath, rel in to_process:
        ext = fpath.suffix.lower()
        label = f"[{section or 'root'}] {rel}"
        print(f"  {label[:78]}", end=" ... ", flush=True)

        if ext in (".doc", ".rtf"):
            texto, error = doc_results.get(str(fpath.resolve()), ("", "doc_com_falhou"))
        else:
            texto, error = extract_generic(fpath)

        entry = {
            "section" : section,
            "arquivo" : rel,
            "tipo"    : ext,
            "texto"   : texto,
            "error"   : error,
        }
        new_entries.append(entry)
        sec_key = section or "root"
        section_counts[sec_key] = section_counts.get(sec_key, 0) + 1
        new_chars += len(texto)

        if error:
            error_counts[error] = error_counts.get(error, 0) + 1
            print(f"ERRO: {error}")
        else:
            print(f"OK ({len(texto):,} chars)")

    # ── Merge e salvar ────────────────────────────────────────────────────────
    all_entries = list(existing.values()) + new_entries

    # Aplicar curadoria
    curadoria_file = _DESP_DIR / "curadoria_corpus.json"
    all_entries, cur_stats = aplicar_curadoria(all_entries, curadoria_file)

    _OUTPUT_FILE.write_text(
        json.dumps(all_entries, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    # ── Resumo final ──────────────────────────────────────────────────────────
    total_chars = sum(len(e.get("texto") or "") for e in all_entries)

    print("\n" + "═" * 62)
    print("RESUMO FINAL")
    print(f"  Total no índice (pós-curadoria) : {len(all_entries)}")
    print(f"  Novos nesta execução            : {len(new_entries)}")
    print(f"  Pulados (já indexados)          : {skipped}")

    if cur_stats.get("excluidos_chunks", 0) > 0 or cur_stats.get("reclassificados_chunks", 0) > 0 or cur_stats.get("rebaixados_chunks", 0) > 0:
        print("\n  Curadoria aplicada:")
        print(f"    Chunks Excluídos              : {cur_stats['excluidos_chunks']} (de {len(cur_stats['excluidos_arquivos'])} arquivos)")
        print(f"    Chunks Reclassificados        : {cur_stats['reclassificados_chunks']} (de {len(cur_stats['reclassificados_arquivos'])} arquivos)")
        print(f"    Chunks Rebaixados             : {cur_stats['rebaixados_chunks']} (de {len(cur_stats['rebaixados_arquivos'])} arquivos)")

    if section_counts:
        print("\n  Novos por seção:")
        for sec, cnt in sorted(section_counts.items()):
            print(f"    {sec:<24}: {cnt}")

    print("\n  Erros:")
    if error_counts:
        for code, cnt in sorted(error_counts.items()):
            print(f"    {code:<38}: {cnt}")
    else:
        print("    (nenhum)")

    print(f"\n  Total de caracteres no índice : {total_chars:,}")
    print(f"  Índice salvo em               : {_OUTPUT_FILE}")
    print("═" * 62)


if __name__ == "__main__":
    main()
